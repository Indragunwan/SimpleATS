"""Extract plain text from uploaded files (PDF/DOCX/TXT)."""
import io


def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF with multi-strategy fallback.

    Priority:
    1. pdfplumber  — best for CV layouts (sorts by Y/X coords, handles multi-column)
    2. pypdf       — fast fallback for simple PDFs
    3. opendataloader_pdf — markdown extraction as last resort
    """
    import io
    import logging
    log = logging.getLogger(__name__)

    # ── Strategy 1: pdfplumber (coordinate-aware, handles multi-column) ──────
    try:
        import pdfplumber

        def _words_to_text(words: list) -> str:
            """Convert sorted word list to plain text, grouping by Y proximity."""
            lines: list[str] = []
            current_line: list[str] = []
            current_top: float = -999.0
            for w in words:
                if abs(w["top"] - current_top) > 5:
                    if current_line:
                        lines.append(" ".join(current_line))
                    current_line = [w["text"]]
                    current_top = w["top"]
                else:
                    current_line.append(w["text"])
            if current_line:
                lines.append(" ".join(current_line))
            return "\n".join(lines)

        def _detect_column_split(words: list, page_width: float) -> float | None:
            """
            Detect if page has two distinct columns by finding a large horizontal gap
            in word x0 positions. Returns the x split point or None if single-column.
            """
            if not words:
                return None
            x0s = sorted(set(round(w["x0"]) for w in words))
            if len(x0s) < 4:
                return None
            # Find largest gap between consecutive x0 clusters
            max_gap = 0
            split_x = None
            for i in range(1, len(x0s)):
                gap = x0s[i] - x0s[i - 1]
                if gap > max_gap:
                    max_gap = gap
                    split_x = (x0s[i - 1] + x0s[i]) / 2
            # Only consider it a two-column layout if:
            # - gap is significant (>40pt) AND
            # - split point is roughly in the middle third of the page (25%–75%)
            if max_gap > 40 and page_width * 0.25 < split_x < page_width * 0.75:
                return split_x
            return None

        pages_text: list[str] = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                words = page.extract_words(
                    x_tolerance=3,
                    y_tolerance=3,
                    keep_blank_chars=False,
                    use_text_flow=False,  # we handle ordering ourselves
                    extra_attrs=["size"],
                )
                if not words:
                    plain = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
                    pages_text.append(plain)
                    continue

                split_x = _detect_column_split(words, page.width)

                if split_x is not None:
                    # Two-column layout: sort each column independently top→bottom
                    # then concatenate: right column first (usually main content / name),
                    # then left column (sidebar: skills, education, contact)
                    right = sorted(
                        [w for w in words if w["x0"] >= split_x],
                        key=lambda w: (round(w["top"] / 5) * 5, w["x0"]),
                    )
                    left = sorted(
                        [w for w in words if w["x0"] < split_x],
                        key=lambda w: (round(w["top"] / 5) * 5, w["x0"]),
                    )
                    # Heuristic: if left column has the candidate name (top-most word
                    # near page top), put left first. Otherwise right first.
                    left_top = min((w["top"] for w in left), default=9999)
                    right_top = min((w["top"] for w in right), default=9999)
                    if left_top < right_top - 10:
                        page_text = _words_to_text(left) + "\n\n" + _words_to_text(right)
                    else:
                        page_text = _words_to_text(right) + "\n\n" + _words_to_text(left)
                else:
                    # Single-column: sort top→bottom, left→right
                    words_sorted = sorted(
                        words, key=lambda w: (round(w["top"] / 5) * 5, w["x0"])
                    )
                    page_text = _words_to_text(words_sorted)

                pages_text.append(page_text)

        result = "\n\n".join(pages_text).strip()
        if result and len(result) > 100:
            return result
        log.warning("pdfplumber returned very short text, trying pypdf")
    except Exception as e:
        log.warning(f"pdfplumber failed: {e}. Falling back to pypdf.")

    # ── Strategy 2: pypdf ────────────────────────────────────────────────────
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        parts: list[str] = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
        result = "\n".join(parts).strip()
        if result and len(result) > 50:
            return result
        log.warning("pypdf returned very short text, trying opendataloader_pdf.")
    except Exception as e:
        log.warning(f"pypdf failed: {e}. Trying opendataloader_pdf.")

    # ── Strategy 3: opendataloader_pdf (markdown) ────────────────────────────
    try:
        import tempfile, os
        import opendataloader_pdf

        with tempfile.TemporaryDirectory() as tmp_dir:
            pdf_path = os.path.join(tmp_dir, "cv.pdf")
            with open(pdf_path, "wb") as f:
                f.write(content)
            out_dir = os.path.join(tmp_dir, "output")
            os.makedirs(out_dir, exist_ok=True)
            opendataloader_pdf.convert(input_path=[pdf_path], output_dir=out_dir, format="markdown")
            for fname in os.listdir(out_dir):
                if fname.endswith(".md"):
                    with open(os.path.join(out_dir, fname), "r", encoding="utf-8") as f_out:
                        text = f_out.read()
                    if text.strip():
                        return text.strip()
    except Exception as e:
        log.warning(f"opendataloader_pdf also failed: {e}")

    return "[Gagal mengekstrak teks dari PDF]"


def extract_text_from_docx(content: bytes) -> str:
    try:
        from docx import Document

        doc = Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(parts).strip()
    except Exception as e:
        return f"[Gagal parsing DOCX: {e}]"


def extract_text(filename: str, content: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return extract_text_from_pdf(content)
    if name.endswith(".docx") or name.endswith(".doc"):
        return extract_text_from_docx(content)
    try:
        return content.decode("utf-8", errors="ignore").strip()
    except Exception:
        return ""
